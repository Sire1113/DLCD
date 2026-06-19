from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

WORKSPACE = Path('.')
RUNS_DIR = WORKSPACE / 'runs'
REPORT_MD = RUNS_DIR / 'report_summary.md'
COMPARISON_JSON = RUNS_DIR / 'comparison.json'
OUT_DOCX = RUNS_DIR / 'dlcd_full_report.docx'

CODE_FILES = [
    Path('src/dlcd/data.py'),
    Path('src/dlcd/models.py'),
    Path('src/dlcd/engine.py'),
]


def _split_markdown_table_row(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip('|').split('|')]


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip('|')
    if not stripped:
        return False
    cells = [cell.strip() for cell in stripped.split('|')]
    return all(cell and set(cell) <= {'-', ':'} for cell in cells)


def _add_markdown_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 2:
        return
    header = _split_markdown_table_row(lines[0])
    body_lines = [line for line in lines[1:] if not _is_table_separator(line)]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    for index, cell_text in enumerate(header):
        table.rows[0].cells[index].text = cell_text
    for line in body_lines:
        row = table.add_row().cells
        for index, cell_text in enumerate(_split_markdown_table_row(line)):
            if index < len(row):
                row[index].text = cell_text


def add_title_page(doc: Document) -> None:
    doc.add_heading('作业二：图像分类任务', level=0)
    p = doc.add_paragraph('基于 ResNet18 的 dhole vs fox 二分类迁移学习实验报告')
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph('')
    doc.add_paragraph('作者：自动生成')
    doc.add_paragraph('日期：2026-06-19')
    doc.add_page_break()


def add_section(doc: Document, title: str, text: str | None = None) -> None:
    doc.add_heading(title, level=1)
    if text:
        for para in text.strip().split('\n\n'):
            doc.add_paragraph(para.strip())


def insert_markdown(doc: Document, md_path: Path) -> None:
    if not md_path.exists():
        doc.add_paragraph('(未找到自动生成的摘要，下面请补充实验结果与分析)')
        return
    text = md_path.read_text(encoding='utf-8')
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            p.add_run(line[2:].strip())
        elif line.startswith('|') and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1].rstrip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith('|'):
                table_lines.append(lines[index].rstrip())
                index += 1
            _add_markdown_table(doc, table_lines)
            continue
        else:
            doc.add_paragraph(line)
        index += 1


def attach_images_for_experiments(doc: Document, comp_json: Path, runs_dir: Path) -> None:
    if not comp_json.exists():
        doc.add_paragraph('(未找到 comparisons.json，无法插入实验图像)')
        return
    comparisons = json.loads(comp_json.read_text(encoding='utf-8'))
    for exp in comparisons:
        name = exp.get('experiment', 'unnamed')
        doc.add_heading(f'实验：{name}', level=2)
        artifact_dir = runs_dir / name
        if not artifact_dir.exists():
            doc.add_paragraph('(该实验没有产出文件夹)')
            continue
        imgs = ['history.png', 'confusion_matrix.png', 'correct_examples.png', 'wrong_examples.png']
        for img in imgs:
            pth = artifact_dir / img
            if pth.exists():
                doc.add_paragraph(img)
                try:
                    doc.add_picture(str(pth), width=Inches(5.5))
                except Exception:
                    doc.add_paragraph(f'无法插入图片 {pth}')
        doc.add_page_break()


def add_code_snippet(doc: Document, path: Path, max_lines: int = 200) -> None:
    doc.add_heading(f'主要代码摘录：{path.name}', level=2)
    if not path.exists():
        doc.add_paragraph('(文件不存在)')
        return
    text = path.read_text(encoding='utf-8')
    lines = text.splitlines()
    excerpt = lines[:max_lines]
    for ln in excerpt:
        p = doc.add_paragraph()
        run = p.add_run(ln)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
    if len(lines) > max_lines:
        doc.add_paragraph('...（截断）')


def build_full_docx():
    doc = Document()
    add_title_page(doc)

    # Insert auto summary if available
    insert_markdown(doc, REPORT_MD)

    doc.add_page_break()

    doc.add_heading('实验图示与结果', level=1)
    attach_images_for_experiments(doc, COMPARISON_JSON, RUNS_DIR)

    # Code appendix (limit to 1-2 pages worth per file by max_lines)
    doc.add_page_break()
    doc.add_heading('附录：主要代码（1-2 页）', level=1)
    for cf in CODE_FILES:
        add_code_snippet(doc, cf, max_lines=28)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(str(OUT_DOCX))


def main() -> None:
    build_full_docx()


if __name__ == '__main__':
    main()
