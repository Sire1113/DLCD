from __future__ import annotations

import json
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


WORKSPACE = Path(".")
RUNS_DIR = WORKSPACE / "runs"
REPORT_MD = RUNS_DIR / "report_summary.md"
COMPARISON_JSON = RUNS_DIR / "comparison.json"
OUT_DOCX = RUNS_DIR / "dlcd_report.docx"

BEST_EXPERIMENT = "finetune_last_block_resnet18"


def _format_center_paragraph(paragraph, *, font_size: int | None = None, bold: bool = False) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if paragraph.runs:
        run = paragraph.runs[0]
    else:
        run = paragraph.add_run()
    run.bold = bold
    if font_size is not None:
        run.font.size = Pt(font_size)


def _add_cover_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.add_run("作业二：图像分类任务报告")
    _format_center_paragraph(title, font_size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.add_run("基于 ResNet18 的 dhole vs fox 二分类迁移学习实验")
    _format_center_paragraph(subtitle, font_size=12)

    doc.add_paragraph("")

    info_lines = [
        "课程：深度学习概论",
        "姓名：________",
        "学号：________",
        "日期：2026-06-19",
    ]
    for line in info_lines:
        p = doc.add_paragraph()
        p.add_run(line)
        _format_center_paragraph(p, font_size=11)

    doc.add_page_break()


def _split_markdown_table_row(line: str) -> list[str]:
    parts = [cell.strip() for cell in line.strip().strip("|").split("|")]
    return parts


def _is_table_separator(line: str) -> bool:
    stripped = line.strip().strip("|")
    if not stripped:
        return False
    cells = [cell.strip() for cell in stripped.split("|")]
    return all(cell and set(cell) <= {"-", ":"} for cell in cells)


def _add_markdown_table(doc: Document, lines: list[str]) -> None:
    if len(lines) < 2:
        return
    header = _split_markdown_table_row(lines[0])
    body_lines = [line for line in lines[1:] if not _is_table_separator(line)]
    table = doc.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for index, cell_text in enumerate(header):
        header_cells[index].text = cell_text
    for line in body_lines:
        row = table.add_row().cells
        for index, cell_text in enumerate(_split_markdown_table_row(line)):
            if index < len(row):
                row[index].text = cell_text


def insert_markdown_as_paragraphs(doc: Document, md_path: Path) -> None:
    if not md_path.exists():
        doc.add_paragraph("(No summary markdown found.)")
        return
    text = md_path.read_text(encoding="utf-8")
    lines = text.splitlines()
    index = 0
    while index < len(lines):
        line = lines[index].rstrip()
        line = line.rstrip()
        if not line:
            doc.add_paragraph("")
            index += 1
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("|") and index + 1 < len(lines) and _is_table_separator(lines[index + 1]):
            table_lines = [line, lines[index + 1].rstrip()]
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                table_lines.append(lines[index].rstrip())
                index += 1
            _add_markdown_table(doc, table_lines)
            continue
        elif line.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(line[2:].strip())
        else:
            doc.add_paragraph(line)
        index += 1


def _get_experiment_artifact_dir(runs_dir: Path, experiment_name: str) -> Path:
    return runs_dir / experiment_name


def _add_image_block(doc: Document, image_path: Path, caption: str, width: float = 5.8) -> None:
    if not image_path.exists():
        return
    caption_paragraph = doc.add_paragraph()
    caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = caption_paragraph.add_run(caption)
    run.font.size = Pt(10)
    run.italic = True
    try:
        doc.add_picture(str(image_path), width=Inches(width))
    except Exception:
        doc.add_paragraph(f"Failed to attach {image_path}")


def attach_interleaved_images(doc: Document, comp_json: Path, runs_dir: Path) -> None:
    if not comp_json.exists():
        return
    comparisons = json.loads(comp_json.read_text(encoding="utf-8"))
    if not comparisons:
        return

    best_artifact_dir = _get_experiment_artifact_dir(runs_dir, BEST_EXPERIMENT)
    linear_artifact_dir = _get_experiment_artifact_dir(runs_dir, "linear_probe_resnet18")
    scratch_artifact_dir = _get_experiment_artifact_dir(runs_dir, "scratch_resnet18")

    doc.add_heading("图示与结果展示", level=1)
    doc.add_paragraph("以下图像穿插在相关章节中，用于对应训练过程、结果分析与错误样例说明。")

    doc.add_heading("训练过程与整体表现", level=2)
    _add_image_block(doc, best_artifact_dir / "history.png", "图 1 预训练微调模型的训练曲线")
    _add_image_block(doc, best_artifact_dir / "confusion_matrix.png", "图 2 预训练微调模型的混淆矩阵")

    doc.add_heading("正确与错误样例", level=2)
    _add_image_block(doc, best_artifact_dir / "correct_examples.png", "图 3 预训练微调模型的正确分类样例")
    _add_image_block(doc, best_artifact_dir / "wrong_examples.png", "图 4 预训练微调模型的错误分类样例")

    doc.add_heading("对比补充", level=2)
    _add_image_block(doc, linear_artifact_dir / "history.png", "图 5 线性探测模型的训练曲线")
    _add_image_block(doc, scratch_artifact_dir / "confusion_matrix.png", "图 6 从头训练模型的混淆矩阵")


def _style_first_paragraph(doc: Document) -> None:
    if not doc.paragraphs:
        return
    first = doc.paragraphs[0]
    first.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if first.runs:
        run = first.runs[0]
        run.bold = True
        run.font.size = Pt(16)


def build_docx():
    doc = Document()
    _add_cover_page(doc)

    insert_markdown_as_paragraphs(doc, REPORT_MD)
    attach_interleaved_images(doc, COMPARISON_JSON, RUNS_DIR)

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)
    print(str(OUT_DOCX))


if __name__ == "__main__":
    build_docx()


def main() -> None:
    """Console entry for project scripts."""
    build_docx()
